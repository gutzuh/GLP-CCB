from flask import Flask, request, send_file, render_template
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

app = Flask(__name__)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/save', methods=['POST'])
def save():
    data = request.get_json()
    musicians = sorted(data.get('musicians', []), key=lambda x: x.lower())
    organists = sorted(data.get('organists', []), key=lambda x: x.lower())
    date_text = data.get('dateText', "03/08/2025 ÁS 17:00H")

    doc = Document()

    # Margens A4
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    usable_width_cm = 21 - 3  # 18 cm

    # Utilidades
    def set_table_layout_fixed(table):
        tbl = table._element
        tblPr = tbl.tblPr
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), str(int(Cm(usable_width_cm).twips)))  # ✅ Corrigido
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)

    def set_repeat_table_header(row):
        tr = row._element
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)

    def prevent_row_split(row):
        tr = row._element
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)

    def add_header(doc, date_text, titulo):
        # Cabeçalho centralizado
        header = doc.add_paragraph("CONGREGAÇÃO CRISTÃ NO BRASIL – BASTOS – SP")
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.runs[0]
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = "Cambria"

        subheader = doc.add_paragraph(f"LISTA DE PRESENÇA – ENSAIO LOCAL – {date_text}")
        subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subheader.runs[0]
        run.font.size = Pt(12)
        run.font.name = "Cambria"

        doc.add_paragraph()  # Espaço

        title = doc.add_paragraph(titulo)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.font.size = Pt(14)
        run.bold = True
        run.font.name = "Cambria"
        # Evita título sozinho no fim da página
        title.paragraph_format.keep_with_next = True

    def fill_table(table, header_titles, rows, extra_rows=20):
        table.style = "Table Grid"
        table.autofit = False
        table.allow_autofit = False
        set_table_layout_fixed(table)

        # Larguras: nome = 11cm, assinatura = 7cm
        col_name_width = Cm(11.0)
        col_sign_width = Cm(7.0)

        # Cabeçalho
        hdr_cells = table.rows[0].cells
        set_repeat_table_header(table.rows[0])
        for i, title in enumerate(header_titles):
            hdr_cells[i].text = title
            hdr_cells[i].width = col_name_width if i == 0 else col_sign_width
            for p in hdr_cells[i].paragraphs:
                run = p.runs[0]
                run.font.size = Pt(12)
                run.bold = True
                run.font.name = "Cambria"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.keep_with_next = True
                p.paragraph_format.keep_together = True

        # Linhas com dados
        for row_text in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = row_text.strip()
            row_cells[1].text = ""
            # Formatação das células
            for idx, cell in enumerate(row_cells):
                cell.width = col_name_width if idx == 0 else col_sign_width
                for p in cell.paragraphs:
                    if not p.runs:
                        p.add_run("")
                    run = p.runs[0]
                    run.font.size = Pt(14)
                    run.font.name = "Cambria"
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
            prevent_row_split(table.rows[-1])

        # Linhas extras vazias
        for _ in range(extra_rows):
            row_cells = table.add_row().cells
            row_cells[0].text = ""
            row_cells[1].text = ""
            for idx, cell in enumerate(row_cells):
                cell.width = col_name_width if idx == 0 else col_sign_width
                for p in cell.paragraphs:
                    if not p.runs:
                        p.add_run("")
                    run = p.runs[0]
                    run.font.size = Pt(14)
                    run.font.name = "Cambria"
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
            prevent_row_split(table.rows[-1])

    # Página 1: Músicos
    add_header(doc, date_text, "MÚSICOS")
    table_m = doc.add_table(rows=1, cols=2)
    fill_table(table_m, ["MÚSICOS - NOME", "ASSINATURA"], musicians, extra_rows=15)

    # Página 2: Organistas
    doc.add_page_break()
    add_header(doc, date_text, "ORGANISTAS")
    table_o = doc.add_table(rows=1, cols=2)
    fill_table(table_o, ["ORGANISTAS - NOME", "ASSINATURA"], organists, extra_rows=15)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="LISTA_ENSAIO_LOCAL_EDITADO.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
 
@app.route('/save_comparecimento', methods=['POST'])
def save_comparecimento():
    data = request.get_json()
    new_date = data.get('dateText', "07/09/2025")

    # Formatar data como DD/MM/YYYY (se vier no formato YYYY-MM-DD)
    if '-' in new_date:
        parts = new_date.split('-')
        if len(parts) == 3:
            year, month, day = parts
            new_date = f"{day}/{month}/{year}"

    # Caminho para o template
    template_path = 'static/Comparecimento ensaio local.docx'

    # Abrir o template
    doc = Document(template_path)

    # Substituir todas as ocorrências da data antiga pela nova
    # O template contém "07/09/2025" — vamos substituir por new_date
    old_date_pattern = "07/09/2025"

    for paragraph in doc.paragraphs:
        if old_date_pattern in paragraph.text:
            paragraph.text = paragraph.text.replace(old_date_pattern, new_date)
            # Aplicar formatação ao texto da data
            for run in paragraph.runs:
                if new_date in run.text:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(18)
                    run.bold = True

    # Também verificar nas tabelas (caso esteja dentro de células)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_date_pattern in paragraph.text:
                        paragraph.text = paragraph.text.replace(old_date_pattern, new_date)
                        # Aplicar formatação ao texto da data
                        for run in paragraph.runs:
                            if new_date in run.text:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(18)
                                run.bold = True

    # Salvar em memória e enviar
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="COMPARCIMENTO_ENSAIO_LOCAL.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == '__main__':
    app.run(debug=True)
